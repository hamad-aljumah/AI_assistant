from fastapi import UploadFile
from sqlalchemy.orm import Session
from app.models.database_models import Document
from app.agents.rag_tool import rag_service
from app.config import get_settings
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import uuid
import logging
from typing import List

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentService:
    """Service for document operations"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len
        )
        os.makedirs(settings.upload_dir, exist_ok=True)
    
    async def upload_document(self, file: UploadFile, db: Session) -> dict:
        """Upload and process document"""
        
        # Generate unique filename
        file_ext = os.path.splitext(file.filename)[1]
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(settings.upload_dir, unique_filename)
        
        # Save file
        content = await file.read()
        with open(file_path, "wb") as f:
            f.write(content)
        
        file_size = len(content)
        
        # Extract text based on file type
        text = self._extract_text(file_path, file_ext)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create metadata for each chunk
        metadatas = [
            {
                "source": file.filename,
                "chunk": i,
                "total_chunks": len(chunks)
            }
            for i in range(len(chunks))
        ]
        
        # Add to vector store
        rag_service.add_documents(chunks, metadatas)
        
        # Save to database
        doc = Document(
            filename=unique_filename,
            original_filename=file.filename,
            file_path=file_path,
            file_size=file_size,
            file_type=file_ext[1:],
            chunk_count=len(chunks)
        )
        db.add(doc)
        db.commit()
        db.refresh(doc)
        
        return {
            "id": doc.id,
            "filename": doc.original_filename,
            "file_size": doc.file_size,
            "file_type": doc.file_type,
            "chunk_count": doc.chunk_count,
            "upload_date": doc.upload_date,
            "message": "Document uploaded successfully"
        }
    
    def _extract_text(self, file_path: str, file_ext: str) -> str:
        """Extract text from file"""
        try:
            if file_ext == ".txt":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif file_ext == ".pdf":
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text()
                return text
            elif file_ext == ".docx":
                from docx import Document
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            elif file_ext == ".md":
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            else:
                return ""
        except Exception as e:
            logger.error(f"Error extracting text: {e}")
            return ""
    
    def list_documents(self, db: Session) -> List[Document]:
        """List all documents"""
        return db.query(Document).order_by(Document.upload_date.desc()).all()
    
    def delete_document(self, document_id: int, db: Session) -> bool:
        """Delete document"""
        doc = db.query(Document).filter(Document.id == document_id).first()
        if doc:
            # Delete file
            if os.path.exists(doc.file_path):
                os.remove(doc.file_path)
            # Delete from database
            db.delete(doc)
            db.commit()
            return True
        return False
